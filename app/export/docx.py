from __future__ import annotations

import base64
import io
import re
import shutil
import subprocess
import tempfile
import sys
import urllib.error
import urllib.request
import json
import zlib
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_BULLET_RE = re.compile(r"^\s*[-*+]\s+(.+?)\s*$")
_NUMBERED_RE = re.compile(r"^\s*\d+[.)]\s+(.+?)\s*$")
_SEPARATOR_CELL_RE = re.compile(r"^:?-{3,}:?$")
_INLINE_TOKEN_RE = re.compile(
    r"(\*\*[^*]+\*\*|__[^_]+__|`[^`]+`|\*[^*]+\*|_[^_]+_)",
)


def _is_table_separator(row: Sequence[str]) -> bool:
    return bool(row) and all(_SEPARATOR_CELL_RE.match(cell.strip()) for cell in row)


def _split_table_row(raw_line: str) -> list[str]:
    line = raw_line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def _add_markdown_runs(paragraph, text: str) -> None:
    cursor = 0
    for match in _INLINE_TOKEN_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])

        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("__") and token.endswith("__"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("_") and token.endswith("_"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)

        cursor = match.end()

    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _add_code_block(document: Document, lines: list[str]) -> None:
    code = "\n".join(lines).rstrip()
    if not code:
        return

    paragraph = document.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def _render_mermaid_png_via_mmdc(code: str) -> bytes | None:
    mmdc_path = shutil.which("mmdc") or shutil.which("mmdc.cmd")
    if mmdc_path is None:
        return None

    with tempfile.NamedTemporaryFile(mode="w", suffix=".mmd", delete=False, encoding="utf-8") as tmp_in:
        tmp_in.write(code)
        tmp_in_path = tmp_in.name

    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_out:
        tmp_out_path = tmp_out.name

    try:
        if sys.platform == "win32":
            cmd = f'""{mmdc_path}" -i "{tmp_in_path}" -o "{tmp_out_path}" --quiet"'
            use_shell = True
        else:
            cmd = [mmdc_path, "-i", tmp_in_path, "-o", tmp_out_path, "--quiet"]
            use_shell = False

        completed = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=30,
            check=False,
            shell=use_shell
        )
        if completed.returncode != 0:
            return None

        output = Path(tmp_out_path).read_bytes()
        return output if output else None
    except Exception:
        return None
    finally:
        Path(tmp_in_path).unlink(missing_ok=True)
        Path(tmp_out_path).unlink(missing_ok=True)

def _render_mermaid_png_via_kroki(code: str) -> bytes | None:
    try:
        compressed = zlib.compress(code.encode("utf-8"), 9)
        encoded = base64.urlsafe_b64encode(compressed).decode("ascii")
        url = f"https://kroki.io/mermaid/png/{encoded}"
        request = urllib.request.Request(url, headers={
            "Accept": "image/png",
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        })
        with urllib.request.urlopen(request, timeout=20) as response:  # noqa: S310
            content_type = response.headers.get("Content-Type", "")
            if "image" not in content_type.lower():
                return None
            payload = response.read()
            return payload if payload else None
    except (urllib.error.URLError, TimeoutError, ValueError):
        return None


def _render_mermaid_png_via_mermaid_ink(code: str) -> bytes | None:
    try:
        payload = json.dumps({"code": code, "mermaid": {"theme": "default"}})
        encoded = base64.urlsafe_b64encode(payload.encode("utf-8")).decode("ascii")
        url = f"https://mermaid.ink/img/{encoded}"
        request = urllib.request.Request(url, headers={
            "Accept": "image/png",
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        })
        with urllib.request.urlopen(request, timeout=20) as response:  # noqa: S310
            content_type = response.headers.get("Content-Type", "")
            if "image" not in content_type.lower():
                return None
            payload = response.read()
            return payload if payload else None
    except (urllib.error.URLError, TimeoutError, ValueError):
        return None


def _render_mermaid_png(code: str) -> bytes | None:
    return _render_mermaid_png_via_mmdc(code) or _render_mermaid_png_via_kroki(code) or _render_mermaid_png_via_mermaid_ink(code)


def _add_mermaid_image(document: Document, code: str) -> bool:
    image_payload = _render_mermaid_png(code)
    if not image_payload:
        return False

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(image_payload), width=Inches(6.4))
    return True


def _add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    header = rows[0]
    body_rows = rows[1:] if len(rows) > 1 else []
    max_cols = max(len(row) for row in rows)
    if max_cols <= 0:
        return

    table = document.add_table(rows=1, cols=max_cols)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for idx in range(max_cols):
        cell_text = header[idx] if idx < len(header) else ""
        paragraph = header_cells[idx].paragraphs[0]
        _add_markdown_runs(paragraph, cell_text)

    for row in body_rows:
        table_row = table.add_row().cells
        for idx in range(max_cols):
            cell_text = row[idx] if idx < len(row) else ""
            paragraph = table_row[idx].paragraphs[0]
            _add_markdown_runs(paragraph, cell_text)


def markdown_to_docx_bytes(
    markdown_text: str,
    *,
    title: str = "SRS",
    author: str = "SRS Generator",
    comments: str = "",
) -> bytes:
    """Convert Markdown text into a DOCX byte payload."""
    document = Document()
    document.core_properties.title = title.strip() or "SRS"
    document.core_properties.author = author.strip() or "SRS Generator"
    document.core_properties.comments = comments.strip()

    lines = markdown_text.splitlines()

    in_code_block = False
    code_buffer: list[str] = []
    code_block_language = ""

    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                if code_block_language == "mermaid":
                    rendered = _add_mermaid_image(document, "\n".join(code_buffer).strip())
                    if not rendered:
                        note = document.add_paragraph("[Diagram image unavailable]")
                        note.runs[0].italic = True
                else:
                    _add_code_block(document, code_buffer)
                code_buffer = []
                code_block_language = ""
                in_code_block = False
            else:
                in_code_block = True
                code_block_language = stripped[3:].strip().lower()
            index += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            index += 1
            continue

        if not stripped:
            document.add_paragraph("")
            index += 1
            continue

        heading_match = _HEADING_RE.match(stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 6)
            text = heading_match.group(2).strip()
            paragraph = document.add_heading(level=level)
            _add_markdown_runs(paragraph, text)
            index += 1
            continue

        if stripped.startswith("|"):
            table_rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_rows.append(_split_table_row(lines[index]))
                index += 1

            normalized_rows: list[list[str]] = []
            for row in table_rows:
                if _is_table_separator(row):
                    continue
                normalized_rows.append(row)

            _add_table(document, normalized_rows)
            continue

        bullet_match = _BULLET_RE.match(line)
        if bullet_match:
            paragraph = document.add_paragraph(style="List Bullet")
            _add_markdown_runs(paragraph, bullet_match.group(1).strip())
            index += 1
            continue

        numbered_match = _NUMBERED_RE.match(line)
        if numbered_match:
            paragraph = document.add_paragraph(style="List Number")
            _add_markdown_runs(paragraph, numbered_match.group(1).strip())
            index += 1
            continue

        if stripped.startswith(">"):
            paragraph = document.add_paragraph()
            _add_markdown_runs(paragraph, stripped.lstrip("> ").strip())
            index += 1
            continue

        paragraph = document.add_paragraph()
        _add_markdown_runs(paragraph, stripped)
        index += 1

    if code_buffer:
        if code_block_language == "mermaid":
            rendered = _add_mermaid_image(document, "\n".join(code_buffer).strip())
            if not rendered:
                note = document.add_paragraph("[Diagram image unavailable]")
                note.runs[0].italic = True
        else:
            _add_code_block(document, code_buffer)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
